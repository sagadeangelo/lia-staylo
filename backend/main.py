# backend/main.py
from __future__ import annotations

import io
import json
import logging
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import requests
import yaml
from fastapi import FastAPI, File, Form, HTTPException, UploadFile, Request, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel

# =========================
# Config & Paths
# =========================
RULES_DIR = Path(os.environ.get("LIA_RULES_DIR", Path(__file__).parent / "rules"))
print("RULES_DIR =", RULES_DIR.resolve())

BASE_DIR = Path(__file__).resolve().parent
STORAGE_DIR = BASE_DIR / "storage"
STORAGE_DIR.mkdir(parents=True, exist_ok=True)
DICT_PATH = STORAGE_DIR / "dictionaries.json"

# LT por defecto en 8010
LT_BASE = (os.environ.get("LT_URL", "http://127.0.0.1:8010") or "").strip().rstrip("/")

def lt_ep(path: str) -> str:
    if not path.startswith("/"):
        path = "/" + path
    return f"{LT_BASE}{path}"

SUPPORTED_UI_LANGS = ("es-MX", "es-419", "en-US")

BRAND_PALETTE = {
    "brand": {"primary": "#FF6A4D", "primaryAlt": "#FF936A", "primarySolid": "#FF7A59", "navy": "#24364B"},
    "light": {
        "background": "#FCFCFD", "surface": "#FFFFFF", "onSurface": "#1F2937", "muted": "#6B7280",
        "divider": "#E5E7EB", "primaryContainer": "#FFE5DE", "onPrimary": "#FFFFFF",
    },
    "dark": {
        "background": "#0E141B", "surface": "#111827", "onSurface": "#DDE3EA",
        "primaryContainer": "#5B2A22", "onPrimary": "#FFFFFF",
    },
    "semantic": {"teal": "#2EC4B6", "info": "#3BA3FF", "success": "#2CB764", "warning": "#FFB020", "error": "#E53935", "outline": "#94A3B8"},
}

logger = logging.getLogger("lia-backend")
logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s")

# =========================
# Helpers: idioma
# =========================
def normalize_lang_ui(lang: str) -> str:
    s = (lang or "").strip()
    if not s:
        return "es-MX"
    s_low = s.lower()
    if s_low.startswith("en"):
        return "en-US"
    if s_low.startswith("es-419"):
        return "es-419"
    return "es-MX"

def pick_lang_ui(param_lang: Optional[str], request: Optional[Request] = None, body_variant: Optional[str] = None) -> str:
    hdr = ""
    try:
        if request is not None:
            hdr = (request.headers.get("X-Lang-Var") or "").strip()
    except Exception:
        hdr = ""
    if hdr in SUPPORTED_UI_LANGS:
        return hdr
    if (body_variant or "") in SUPPORTED_UI_LANGS:
        return str(body_variant)
    return normalize_lang_ui(param_lang or "")

def to_lt_language(lang_ui: str) -> str:
    return "en-US" if lang_ui.startswith("en") else "es"

# =========================
# Diccionario persistente
# =========================
def _ensure_dict_file() -> Dict[str, List[str]]:
    if not DICT_PATH.exists():
        data = {k: [] for k in SUPPORTED_UI_LANGS}
        DICT_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return data
    try:
        return json.loads(DICT_PATH.read_text(encoding="utf-8"))
    except Exception:
        data = {k: [] for k in SUPPORTED_UI_LANGS}
        DICT_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return data

def _save_dict_file(data: Dict[str, List[str]]) -> None:
    DICT_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def dict_list(lang_ui: str) -> List[str]:
    data = _ensure_dict_file()
    return sorted(set(map(str, data.get(lang_ui, []))))

def dict_add(lang_ui: str, token: str) -> None:
    token = token.strip()
    if not token:
        raise ValueError("Token vacío.")
    if " " in token:
        raise ValueError("El diccionario solo acepta UNA palabra (sin espacios).")
    data = _ensure_dict_file()
    arr = set(map(str, data.get(lang_ui, [])))
    arr.add(token)
    data[lang_ui] = sorted(arr)
    _save_dict_file(data)

def dict_remove(lang_ui: str, token: str) -> None:
    token = token.strip()
    data = _ensure_dict_file()
    arr = set(map(str, data.get(lang_ui, [])))
    if token in arr:
        arr.remove(token)
        data[lang_ui] = sorted(arr)
        _save_dict_file(data)

# =========================
# Reglas personalizadas (YAML sencillo)
# =========================
@dataclass
class CustomRule:
    id: str
    message: str
    short: str
    regex: re.Pattern
    category: str = "STYLE"
    suggestions: Tuple[str, ...] = tuple()

_RULES_CACHE: Dict[str, List[CustomRule]] = {}

def _compile_rule(raw: dict, fallback_prefix: str) -> Optional[CustomRule]:
    rid = str(raw.get("id") or raw.get("rule") or f"{fallback_prefix}_RULE")
    message = str(raw.get("message") or raw.get("msg") or "Observación de estilo.")
    short = str(raw.get("shortMessage") or raw.get("short") or "")
    category = str(raw.get("category") or "STYLE").upper()
    pat = raw.get("regex") or raw.get("pattern")
    if not pat and isinstance(raw.get("tokens"), list):
        toks = [re.escape(str(t)) for t in raw["tokens"] if str(t).strip()]
        if toks:
            pat = r"\b(?:%s)\b" % "|".join(toks)
    if not pat:
        return None
    try:
        rgx = re.compile(str(pat), re.IGNORECASE | re.MULTILINE)
    except re.error:
        return None
    sugg = tuple(map(str, raw.get("suggestions") or raw.get("replacements") or []))
    return CustomRule(id=rid, message=message, short=short, regex=rgx, category=category, suggestions=sugg)

def _load_yaml_rules(path: Path) -> List[dict]:
    raw = yaml.safe_load(path.read_text(encoding="utf-8"))
    if raw is None:
        return []
    if isinstance(raw, dict):
        if "rules" in raw and isinstance(raw["rules"], list):
            return raw["rules"]
        if "items" in raw and isinstance(raw["items"], list):
            return raw["items"]
        return [v for v in raw.values() if isinstance(v, dict)]
    if isinstance(raw, list):
        return raw
    return []

def load_custom_rules(lang_ui: str) -> List[CustomRule]:
    if lang_ui in _RULES_CACHE:
        return _RULES_CACHE[lang_ui]
    fname = "es_mx.yaml" if lang_ui == "es-MX" else "es_419.yaml" if lang_ui == "es-419" else None
    rules: List[CustomRule] = []
    if fname:
        path = RULES_DIR / fname
        if path.exists():
            try:
                items = _load_yaml_rules(path)
                for i, r in enumerate(items):
                    if isinstance(r, dict):
                        obj = _compile_rule(r, f"CUSTOM_{fname}_{i}")
                        if obj:
                            rules.append(obj)
            except Exception as e:
                logger.warning("No se pudieron cargar reglas %s: %s", path, e)
    _RULES_CACHE[lang_ui] = rules
    logger.info("Reglas cargadas %s: %d", lang_ui, len(rules))
    return rules

def run_custom_rules(text: str, lang_ui: str) -> List[dict]:
    out: List[dict] = []
    for rule in load_custom_rules(lang_ui):
        for m in rule.regex.finditer(text):
            start, end = m.start(), m.end()
            out.append(
                {
                    "message": rule.message,
                    "shortMessage": rule.short,
                    "offset": start,
                    "length": max(0, end - start),
                    "replacements": [{"value": v} for v in list(rule.suggestions)],
                    "rule": {
                        "id": rule.id,
                        "description": rule.message,
                        "issueType": "style",
                        "category": {"id": "STYLE", "name": "Style"},
                        "isPremium": False,
                    },
                    "clientClass": "style",
                    "lt_clientClass": "style",
                }
            )
    return out

# =========================
# LT client + clasificación robusta
# =========================
def _classify_client(m: dict) -> str:
    rule = m.get("rule") or {}
    rid = str(rule.get("id") or "").upper()
    cat = rule.get("category") or {}
    cid = str((cat.get("id") or cat.get("name") or "")).upper()
    if any(k in rid for k in ("SPELL", "MORFOLOGIK", "HUNSPELL", "MISSPELL", "TYPO")) or any(k in cid for k in ("SPELL", "TYPOS", "TYPO")):
        return "spelling"
    if any(k in rid for k in ("PUNCT", "COMMA", "QUOTES", "DASH", "ELLIPSIS", "WHITESPACE", "SPACE")) or \
       any(k in cid for k in ("PUNCT", "COMMA", "QUOTES", "DASH", "ELLIPSIS", "WHITESPACE", "SPACE", "TYPOGRAPHY")):
        return "punct"
    issue = str(rule.get("issueType") or "").lower()
    if "style" in issue or "STYLE" in cid:
        return "style"
    return "grammar"

def lt_check(text: str, lang_ui: str, timeout_s: int = 30) -> List[dict]:
    lt_lang = to_lt_language(lang_ui)
    url = lt_ep("/v2/check")
    try:
        resp = requests.post(url, data={"language": lt_lang, "text": text}, timeout=timeout_s)
        resp.raise_for_status()
    except Exception as e:
        raise RuntimeError(f"LanguageTool no disponible en {LT_BASE}: {e}") from e
    data = resp.json()
    matches = data.get("matches", [])
    out = []
    for m in matches:
        cls = _classify_client(m)
        m["clientClass"] = cls
        m["lt_clientClass"] = cls
        out.append(m)
    return out

# =========================
# Stats / Readability (ligero)
# =========================
_SENT_SPLIT = re.compile(r"[\.!\?…;:\n\r]+")
_VOWELS = re.compile(r"[aeiouáéíóúü]+", re.I)

def basic_stats(text: str) -> Dict[str, int]:
    words = re.findall(r"\w+", text, flags=re.UNICODE)
    sentences = [s.strip() for s in _SENT_SPLIT.split(text) if s.strip()]
    long_sentences = sum(1 for s in sentences if len(s.split()) > 30)
    dialog_marks = text.count("—") + text.count("―") + text.count("“") + text.count("”")
    return {"words": len(words), "sentences": len(sentences), "long_sentences": int(long_sentences), "dialog_marks": int(dialog_marks)}

def rough_syllables(text: str) -> int:
    return len(_VOWELS.findall(text))

def readability_info(text: str) -> Dict[str, float]:
    w = max(1, len(re.findall(r"\w+", text)))
    s = max(1, len([t for t in _SENT_SPLIT.split(text) if t.strip()]))
    y = max(1, rough_syllables(text))
    flesch = 206.835 - 1.015 * (w / s) - 84.6 * (y / w)
    return {"flesch_en_reference": float(round(flesch, 6)), "syllables": int(y)}

# =========================
# Filtrado ortografía por diccionario de usuario
# =========================
def _is_spelling_match(m: dict) -> bool:
    cls = str(m.get("clientClass") or m.get("lt_clientClass") or "")
    return cls == "spelling"

def filter_spelling_by_user_dict(matches: List[dict], text: str, lang_ui: str) -> List[dict]:
    user_words = set(map(str.lower, dict_list(lang_ui)))
    if not user_words:
        return matches
    out = []
    for m in matches:
        if _is_spelling_match(m):
            off = int(m.get("offset") or 0)
            ln = int(m.get("length") or 0)
            token = text[off: off + ln].lower()
            if token in user_words:
                continue
        out.append(m)
    return out

# =========================
# FastAPI app
# =========================
app = FastAPI(title="LIA-Staylo API", version="0.8.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_headers=["*"], allow_methods=["*"])

# -------- Models --------
class AnalyzeTextIn(BaseModel):
    text: str
    lang: Optional[str] = "es-MX"
    variant: Optional[str] = None

class ApplyIn(BaseModel):
    text: str
    lang: Optional[str] = "es-MX"
    variant: Optional[str] = None

class DictIn(BaseModel):
    token: str
    lang: Optional[str] = "es-MX"
    variant: Optional[str] = None

class SuggestIn(BaseModel):
    text: str
    lang: Optional[str] = "es-MX"
    variant: Optional[str] = None

# -------- Health --------
@app.get("/health")
def health():
    ok = True
    lt_ok = False
    try:
        r = requests.get(lt_ep("/v2/languages"), timeout=5)
        lt_ok = r.ok
    except Exception:
        lt_ok = False
    return {
        "ok": ok,
        "lt_ok": lt_ok,
        "ltOk": lt_ok,
        "lt_url": LT_BASE,
        "langs": list(SUPPORTED_UI_LANGS),
        "rules": {lg: len(load_custom_rules(lg)) for lg in SUPPORTED_UI_LANGS},
    }

# -------- Brand --------
@app.get("/brand")
def brand():
    return {"brand": BRAND_PALETTE, "version": app.version}

# -------- Analyze: text --------
@app.post("/analyze_text")
def analyze_text(payload: AnalyzeTextIn, request: Request):
    lang_ui = pick_lang_ui(payload.lang, request, payload.variant)
    text = payload.text or ""
    try:
        lt_matches = lt_check(text, lang_ui)               # LT
        custom_matches = run_custom_rules(text, lang_ui)   # Reglas
        matches = lt_matches + custom_matches
        matches = filter_spelling_by_user_dict(matches, text, lang_ui)
    except Exception as e:
        logger.exception("Error analizando texto")
        raise HTTPException(status_code=500, detail=str(e))
    return {
        "ok": True,
        "text": text,
        "language": lang_ui,
        "stats": basic_stats(text),
        "readability": readability_info(text),
        "languageTool": {"matches": matches},
    }

# -------- Analyze: compat JSON con Flutter (/analyze) --------
@app.post("/analyze")
def analyze_compat(payload: dict, request: Request):
    text = str(payload.get("text") or "")
    lang = str(payload.get("lang") or "es-MX")
    variant = payload.get("variant")
    lang_ui = pick_lang_ui(lang, request, variant)
    lt_lang = str(payload.get("ltLang") or to_lt_language(lang_ui))
    res = analyze_text(AnalyzeTextIn(text=text, lang=lang_ui, variant=variant), request)  # type: ignore
    return {
        "input": {"lang": lang_ui, "ltLang": lt_lang, "length": len(text)},
        "ok": True,
        "stats": res["stats"],
        "readability": res["readability"],
        "languageTool": res["languageTool"],
    }

# -------- Analyze: file (txt/md/docx/pdf) --------
def _read_upload_text(upload: UploadFile) -> str:
    name = (upload.filename or "file").lower()
    raw = upload.file.read()
    if name.endswith(".txt") or name.endswith(".md"):
        return raw.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        try:
            from docx import Document  # python-docx
        except Exception as e:
            raise RuntimeError("Falta dependencia 'python-docx'. Instala con: pip install python-docx") from e
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(raw); tmp.flush(); path = tmp.name
        try:
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        finally:
            try: os.unlink(path)
            except Exception: pass
    if name.endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader
        except Exception as e:
            raise RuntimeError("Falta dependencia 'PyPDF2'. Instala con: pip install PyPDF2") from e
        bio = io.BytesIO(raw)
        reader = PdfReader(bio)
        pages = []
        for p in reader.pages:
            pages.append(p.extract_text() or "")
        return "\n".join(pages)
    raise ValueError("Extensión no soportada. Usa .txt, .md, .docx o .pdf")

@app.post("/analyze/file")
def analyze_file(
    file: UploadFile = File(None),
    manuscript: UploadFile = File(None),
    lang: str = Form("es-MX"),
    request: Request = None,
):
    upload = file or manuscript
    if upload is None:
        raise HTTPException(status_code=400, detail="Falta el archivo: use campo 'file' (o 'manuscript').")
    try:
        fname = upload.filename or "(sin nombre)"
        pos = upload.file.seek(0, 2)
        upload.file.seek(0)
        print(f">> /analyze/file recibido: {fname} (bytes ~ {pos})")
    except Exception:
        pass
    lang_ui = pick_lang_ui(lang, request, None)
    try:
        text = _read_upload_text(upload)
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        logger.exception("No se pudo leer el archivo")
        raise HTTPException(status_code=400, detail=f"No se pudo leer el archivo: {e}")
    print(f">> Texto extraído: {len(text)} chars")
    try:
        return analyze_text(AnalyzeTextIn(text=text, lang=lang_ui), request)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Fallo analizando el archivo")
        raise HTTPException(status_code=500, detail=f"Fallo del analizador: {e}")

@app.post("/upload")
@app.post("/api/upload")
def upload_legacy(
    file: UploadFile = File(None),
    manuscript: UploadFile = File(None),
    lang: str = Form("es-MX"),
    request: Request = None,
):
    return analyze_file(file=file, manuscript=manuscript, lang=lang, request=request)

@app.post("/analyze-file")
def analyze_file_alias(
    file: UploadFile = File(None),
    manuscript: UploadFile = File(None),
    lang: str = Form("es-MX"),
    request: Request = None,
):
    return analyze_file(file=file, manuscript=manuscript, lang=lang, request=request)

# -------- Apply: safe / all --------
def _apply_from_matches(text: str, matches: List[dict]) -> str:
    matches_sorted = sorted(matches, key=lambda m: int(m.get("offset") or 0), reverse=True)
    new_text = text
    for m in matches_sorted:
        reps = m.get("replacements") or []
        if not reps:
            continue
        off = int(m.get("offset") or 0)
        ln = int(m.get("length") or 0)
        if off < 0 or ln <= 0 or off + ln > len(new_text):
            continue
        first = reps[0]
        repl = first.get("value") if isinstance(first, dict) and "value" in first else str(first)
        new_text = new_text[:off] + repl + new_text[off + ln:]
    return new_text

@app.post("/apply/safe")
def apply_safe(payload: ApplyIn, request: Request):
    lang_ui = pick_lang_ui(payload.lang, request, payload.variant)
    text = payload.text or ""
    lt_matches = lt_check(text, lang_ui)
    safe = []
    for m in lt_matches:
        rule = m.get("rule") or {}
        rid = str(rule.get("id") or "").upper()
        cid = str(((rule.get("category") or {}).get("id") or "")).upper()
        if any(k in rid for k in ("COMMA", "WHITESPACE", "PUNCT", "ELLIPSIS", "DASH", "APOS")) or any(k in cid for k in ("PUNCT", "WHITESPACE")):
            safe.append(m)
    new_text = _apply_from_matches(text, safe)
    return {"new_text": new_text}

@app.post("/apply/all")
def apply_all(payload: ApplyIn, request: Request):
    lang_ui = pick_lang_ui(payload.lang, request, payload.variant)
    text = payload.text or ""
    lt_matches = lt_check(text, lang_ui)
    new_text = _apply_from_matches(text, lt_matches)
    return {"new_text": new_text}

# -------- Fallback helpers (para /suggest)
_SMS_MAP = [
    (r"\bxq\b", "porque"),
    (r"\bq\b", "que"),
    (r"\bke\b", "que"),
    (r"\bk\b", "que"),
    (r"\bpa\b", "para"),
    (r"\bxa\b", "para"),
    (r"\bxfa\b", "por favor"),
    (r"\bxo\b", "pero"),
    (r"\bd\b", "de"),
    (r"\bt\b", "te"),
    (r"\btoi\b", "estoy"),
    (r"\btoy\b", "estoy"),
    (r"\bkpasa\b", "qué pasa"),
    (r"\bna\b", "nada"),
    (r"\bpk\b", "porque"),
    (r"\bporq\b", "porque"),
    (r"\bpor que\b", "porque"),
    (r"\bqno\b", "que no"),
    (r"\bq si\b", "que sí"),
    (r"\bkien(es)?\b", "quien"),
    (r"\bvaia\b", "vaya"),
]

_COMMON_FIXES = {
    "cuanras": "cuantas",
    "kuantas": "cuantas",
    "kien": "quien",
    "kienas": "quienes",
    "ai": "hay",
    "ase": "hace",
    "porsupuesto": "por supuesto",
    "averca": "acerca",
    "haver": "haber",
    "aver": "a ver",
    "haiga": "haya",
}

def _fallback_sms_normalize(text: str) -> str:
    s = " " + text + " "
    for pat, repl in _SMS_MAP:
        s = re.sub(pat, repl, s, flags=re.IGNORECASE)
    return s.strip()

def _light_rewrite(text: str) -> str:
    s = re.sub(r"[ \t]+", " ", text)
    s = re.sub(r"\s+\n", "\n", s)
    s = re.sub(r" ,", ",", s)
    s = re.sub(r" \.", ".", s)
    s = re.sub(r"\s{2,}", " ", s).strip()
    return s

# -------- Suggest (reescritura con fallback local)
# -------- Suggest (reescritura con/fallback) --------
@app.post("/suggest")
def suggest(payload: SuggestIn, request: Request):
    """
    Reescritura:
    - por defecto hace 'fix' (usa LT + aplica replacements)
    - modo 'clean' mantiene la limpieza heurística previa
    """
    t = (payload.text or "").strip()
    if not t:
        return {"suggestion": ""}

    mode = (request.query_params.get("mode") or "fix").lower().strip()
    lang_ui = pick_lang_ui(payload.lang, request, payload.variant)

    if mode == "clean":
        s = re.sub(r"[ \t]+", " ", t)
        s = re.sub(r" ,", ",", s)
        s = re.sub(r" \.", ".", s)
        s = re.sub(r"\s+\n", "\n", s)
        s = re.sub(r"\b(pues|este|o\s+sea)\b", "", s, flags=re.I)
        s = re.sub(r"\s{2,}", " ", s).strip()
        return {"suggestion": s}

    # --- modo 'fix' (DEFAULT): usa LanguageTool y aplica replacements ---
    try:
        matches = lt_check(t, lang_ui)
        logger.info("[/suggest] mode=fix lang=%s matches=%d", lang_ui, len(matches))
        fixed = _apply_from_matches(t, matches)
        if fixed == t:
            s = re.sub(r"[ \t]+", " ", t)
            s = re.sub(r"\s{2,}", " ", s).strip()
            return {"suggestion": s}
        return {"suggestion": fixed}
    except Exception as e:
        logger.warning("SUGGEST fallback clean por error LT: %s", e)
        s = re.sub(r"[ \t]+", " ", t)
        s = re.sub(r"\s{2,}", " ", s).strip()
        return {"suggestion": s}


# --- Plan B explícito para probar sin querystring ni heurística ---
@app.post("/suggest_fix_simple")
def suggest_fix_simple(payload: SuggestIn, request: Request):
    """
    Aplica SIEMPRE todos los replacements de LT (sin limpieza).
    Úsalo para verificar que LT devuelve matches y que el apply funciona.
    """
    t = (payload.text or "").strip()
    if not t:
        return {"suggestion": ""}

    lang_ui = pick_lang_ui(payload.lang, request, payload.variant)
    try:
        matches = lt_check(t, lang_ui)
        logger.info("[/suggest_fix_simple] lang=%s matches=%d", lang_ui, len(matches))
        fixed = _apply_from_matches(t, matches)
        return {"suggestion": fixed}
    except Exception as e:
        logger.exception("suggest_fix_simple LT error")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/dictionary/list")
def dictionary_list(lang: str = "es-MX", request: Request = None):
    lang_ui = pick_lang_ui(lang, request, None)
    return {"words": dict_list(lang_ui)}

@app.post("/dictionary/add")
def dictionary_add(payload: DictIn, request: Request):
    lang_ui = pick_lang_ui(payload.lang, request, payload.variant)
    token = (payload.token or "").strip()
    try:
        dict_add(lang_ui, token)
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/dictionary/remove")
def dictionary_remove(payload: DictIn, request: Request):
    lang_ui = pick_lang_ui(payload.lang, request, payload.variant)
    token = (payload.token or "").strip()
    try:
        dict_remove(lang_ui, token)
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# -------- Export: DOCX --------
@app.post("/export/docx")
def export_docx(payload: ApplyIn):
    try:
        from docx import Document
    except Exception as e:
        raise HTTPException(status_code=500, detail="Falta dependencia 'python-docx'. Instala con: pip install python-docx") from e
    doc = Document()
    for line in (payload.text or "").splitlines():
        doc.add_paragraph(line)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name); tmp.seek(0); bin_data = tmp.read()
    return Response(
        content=bin_data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="LIA-Staylo.docx"'},
    )

# -------- Admin: recargar reglas --------
admin = APIRouter()

@admin.post("/admin/reload-rules")
def reload_rules():
    _RULES_CACHE.clear()
    counts = {lg: len(load_custom_rules(lg)) for lg in SUPPORTED_UI_LANGS}
    return {"reloaded": True, "rules": counts}

app.include_router(admin)

# -------- Root --------
@app.get("/")
def root():
    return {"app": "LIA-Staylo API", "version": app.version}

# -------- Run (dev) --------
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=3000, reload=True, log_level="info")
